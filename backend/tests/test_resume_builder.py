import io
from docx import Document
from services.resume_builder import apply_edits_to_docx, _get_all_paragraphs


def _make_docx(lines: list[str]) -> bytes:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _all_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in _get_all_paragraphs(doc))


_RESUME = [
    "PROFESSIONAL SUMMARY",
    "Experienced engineer.",
    "TECHNICAL SKILLS",
    "Python, SQL",
    "PROFESSIONAL EXPERIENCE",
    "Engineer | Acme Corp   2020-2022",
    "● Built internal tools",
]

_EDITS = {
    "updated_summary": "Senior engineer fluent in Kubernetes and Docker.",
    "skills_to_add": ["Kubernetes", "Docker"],
    "experience_enhancements": [
        {"company": "Acme Corp", "title": "Engineer", "new_bullets": ["Shipped feature X", "Owned service Y"]}
    ],
}


def test_all_sections_applied():
    out, report = apply_edits_to_docx(_make_docx(_RESUME), _EDITS)
    assert report["any_applied"] is True
    assert report["summary_updated"] is True
    assert report["skills_added"] is True
    assert report["jobs_enhanced"] == 1

    text = _all_text(out)
    assert "Senior engineer fluent in Kubernetes" in text
    assert "Kubernetes" in text and "Docker" in text
    assert "Shipped feature X" in text
    assert "Owned service Y" in text


def test_original_skills_preserved():
    out, _ = apply_edits_to_docx(_make_docx(_RESUME), _EDITS)
    assert "Python, SQL" in _all_text(out)


def test_unrecognized_format_reports_no_op():
    # No standard headings -> heuristics match nothing -> any_applied False.
    junk = _make_docx(["my resume", "i did stuff", "at a place"])
    out, report = apply_edits_to_docx(junk, _EDITS)
    assert report["any_applied"] is False
    assert report["jobs_enhanced"] == 0


def test_empty_edits_is_no_op():
    out, report = apply_edits_to_docx(_make_docx(_RESUME), {})
    assert report["any_applied"] is False


def _para_by_text(docx_bytes: bytes, needle: str):
    doc = Document(io.BytesIO(docx_bytes))
    for p in _get_all_paragraphs(doc):
        if needle in p.text:
            return p
    return None


def test_skills_categorized_placement_and_dedup():
    doc = Document()
    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph("Cloud & Infrastructure: AWS, Docker")
    doc.add_paragraph("Databases & Tools: MySQL")
    buf = io.BytesIO(); doc.save(buf)

    edits = {"skills_to_add": [
        {"category": "Cloud & Infrastructure", "skills": ["Kubernetes", "Docker"]},
        {"category": "Databases & Tools", "skills": ["PostgreSQL"]},
    ]}
    out, report = apply_edits_to_docx(buf.getvalue(), edits)

    assert report["skills_added"] is True
    text = _all_text(out)
    cloud = _para_by_text(out, "Cloud & Infrastructure").text
    dbs = _para_by_text(out, "Databases & Tools").text
    assert "Kubernetes" in cloud            # new skill placed in matching category
    assert "PostgreSQL" in dbs
    assert text.count("Docker") == 1        # already present -> not duplicated
    assert "Kubernetes" not in dbs          # not dumped onto the last line


def test_list_style_bullet_gets_no_literal_glyph():
    doc = Document()
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Engineer | Acme Corp | 2020")
    doc.add_paragraph("Built dashboards", style="List Bullet")
    buf = io.BytesIO(); doc.save(buf)

    edits = {"experience_enhancements": [
        {"company": "Acme Corp", "title": "Engineer", "new_bullets": ["Automated reporting"]}
    ]}
    out, report = apply_edits_to_docx(buf.getvalue(), edits)

    assert report["jobs_enhanced"] == 1
    p = _para_by_text(out, "Automated reporting")
    assert p is not None
    assert "●" not in p.text                # auto-bulleted list -> no double glyph


def test_inserted_bullet_is_not_bold():
    doc = Document()
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Engineer | Acme Corp | 2020")
    p = doc.add_paragraph()
    p.add_run("● ")
    rb = p.add_run("Project: "); rb.bold = True
    p.add_run("shipped it")
    buf = io.BytesIO(); doc.save(buf)

    edits = {"experience_enhancements": [
        {"company": "Acme Corp", "title": "Engineer", "new_bullets": ["Owned service Y"]}
    ]}
    out, _ = apply_edits_to_docx(buf.getvalue(), edits)

    new_p = _para_by_text(out, "Owned service Y")
    assert new_p is not None
    assert all(r.bold is not True for r in new_p.runs)   # body not forced bold


def test_company_match_tolerates_suffix_difference():
    # Resume has full legal name; LLM returns the short form.
    doc = Document()
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("AI Engineer | Spirit AI Solutions Pvt Ltd | Oct 2025")
    doc.add_paragraph("Built automation pipelines", style="List Bullet")
    buf = io.BytesIO(); doc.save(buf)

    edits = {"experience_enhancements": [
        {"company": "Spirit AI Solutions", "title": "AI Engineer", "new_bullets": ["Integrated LLM workflows"]}
    ]}
    out, report = apply_edits_to_docx(buf.getvalue(), edits)

    assert report["jobs_enhanced"] == 1
    assert "Integrated LLM workflows" in _all_text(out)


def test_job_without_real_bullets_is_skipped():
    # Job title only, no bullets -> must NOT inject a bold line at the title.
    doc = Document()
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Engineer | Acme Corp | 2020")
    buf = io.BytesIO(); doc.save(buf)

    edits = {"experience_enhancements": [
        {"company": "Acme Corp", "title": "Engineer", "new_bullets": ["Should not appear"]}
    ]}
    out, report = apply_edits_to_docx(buf.getvalue(), edits)

    assert report["jobs_enhanced"] == 0
    assert "Should not appear" not in _all_text(out)
