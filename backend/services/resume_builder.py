import io
import re
import copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _para_text(para) -> str:
    return para.text.strip()


def _get_all_paragraphs(doc):
    """Return all paragraphs in document order, including table cells."""
    from docx.text.paragraph import Paragraph
    result = []

    def _collect(element):
        for child in element:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                result.append(Paragraph(child, doc))
            elif tag in ("tbl", "tc", "tr"):
                _collect(child)

    _collect(doc.element.body)
    return result


def _is_section_heading(para) -> bool:
    text = _para_text(para)
    if not text or len(text) < 2:
        return False
    return text.upper() == text and len(text.split()) <= 6


def _find_section(paragraphs, *names):
    """Return index of section heading matching any of the given names."""
    for i, p in enumerate(paragraphs):
        t = _para_text(p).upper()
        for name in names:
            if name.upper() in t:
                return i
    return None


def _section_content_range(paragraphs, heading_idx):
    """Return (start, end) indices of content paragraphs within a section."""
    start = heading_idx + 1
    for i in range(start, len(paragraphs)):
        if _para_text(paragraphs[i]) and _is_section_heading(paragraphs[i]):
            return start, i
    return start, len(paragraphs)


def _run_is_bold(r) -> bool:
    rpr = r.find(f"{{{W}}}rPr")
    if rpr is None:
        return False
    b = rpr.find(f"{{{W}}}b")
    if b is None:
        return False
    return b.get(f"{{{W}}}val") not in ("0", "false", "none")


def _body_rpr(ref_para):
    """Run properties representing normal bullet *body* text.

    Existing bullets often start with a bold lead-in (project title) followed by
    normal-weight description text. Copying the first run's format makes the whole
    new bullet bold, so prefer a non-bold run with text; otherwise strip bold from
    the first run's format.
    """
    runs = ref_para._element.findall(f"{{{W}}}r")
    for r in runs:
        t = r.find(f"{{{W}}}t")
        if t is not None and (t.text or "").strip() and not _run_is_bold(r):
            rpr = r.find(f"{{{W}}}rPr")
            return copy.deepcopy(rpr) if rpr is not None else None
    for r in runs:
        rpr = r.find(f"{{{W}}}rPr")
        if rpr is not None:
            rpr = copy.deepcopy(rpr)
            for b in rpr.findall(f"{{{W}}}b") + rpr.findall(f"{{{W}}}bCs"):
                rpr.remove(b)
            return rpr
    return None


def _para_has_numbering(para) -> bool:
    """True if the paragraph is part of an auto-numbered/bulleted Word list."""
    ppr = para._element.find(f"{{{W}}}pPr")
    return ppr is not None and ppr.find(f"{{{W}}}numPr") is not None


def _leading_glyph(text: str):
    """Return the literal bullet glyph a line starts with, if any."""
    stripped = text.lstrip()
    for g in ("●", "•", "▪", "◦", "‣", "·", "-", "*"):
        if stripped.startswith(g):
            return g
    return None


def _style_is_list(para) -> bool:
    try:
        return bool(para.style and para.style.name and para.style.name.startswith("List"))
    except Exception:
        return False


def _auto_bulleted(para) -> bool:
    """True if Word renders the bullet glyph automatically (numbering or List style)."""
    return _para_has_numbering(para) or _style_is_list(para)


def _is_bullet_para(para) -> bool:
    """A real list item: auto-numbered, a List style, or starting with a glyph."""
    return _auto_bulleted(para) or _leading_glyph(_para_text(para)) is not None


def _insert_para_after(ref_para, text: str):
    """Insert a new paragraph after ref_para, mirroring its bullet style.

    Clones ref_para (keeping paragraph properties such as list numbering and
    indentation), then writes one run using non-bold body formatting so the new
    line matches the existing bullets instead of a bold, mismatched paragraph.
    """
    new_p = copy.deepcopy(ref_para._element)

    # Remove all runs from copy
    for r in new_p.findall(f"{{{W}}}r"):
        new_p.remove(r)
    for r in new_p.findall(f".//{{{W}}}r"):
        parent = r.getparent()
        if parent is not None:
            parent.remove(r)

    rpr_elem = _body_rpr(ref_para)

    r_elem = etree.SubElement(new_p, f"{{{W}}}r")
    if rpr_elem is not None:
        r_elem.insert(0, rpr_elem)
    t_elem = etree.SubElement(r_elem, f"{{{W}}}t")
    t_elem.text = text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    ref_para._element.addnext(new_p)


def _update_summary(paragraphs, new_text: str) -> bool:
    """Replace profile summary paragraph text. Returns True if applied."""
    heading_idx = _find_section(paragraphs, "PROFILE SUMMARY", "SUMMARY", "PROFESSIONAL SUMMARY")
    if heading_idx is None:
        return False

    start, end = _section_content_range(paragraphs, heading_idx)
    for i in range(start, end):
        p = paragraphs[i]
        if _para_text(p):
            # Clear all runs and set new text in first run
            runs = p.runs
            if runs:
                runs[0].text = new_text
                for run in runs[1:]:
                    run.text = ""
            else:
                p.add_run(new_text)
            return True
    return False


def _append_to_para(para, skills: list) -> None:
    """Append a comma-joined list of skills onto a paragraph's last run."""
    runs = para.runs
    appended = ", ".join(skills)
    if runs:
        runs[-1].text = runs[-1].text.rstrip().rstrip(",") + ", " + appended
    else:
        para.add_run(", " + appended)


def _label_of(text: str) -> str:
    """Category label of a skills line, e.g. 'Cloud & Infrastructure' from 'Cloud & Infrastructure: ...'."""
    return text.split(":", 1)[0] if ":" in text else ""


def _best_category_para(skill_paras, category: str):
    """Pick the skills line whose label best matches a category, by word overlap."""
    want = {w for w in re.findall(r"[a-z0-9]+", category.lower()) if len(w) > 2}
    best, best_score = None, 0
    for p in skill_paras:
        label = _label_of(_para_text(p))
        if not label:
            continue
        have = {w for w in re.findall(r"[a-z0-9]+", label.lower()) if len(w) > 2}
        score = len(want & have)
        if score > best_score:
            best, best_score = p, score
    return best


def _add_skills(paragraphs, skills: list) -> bool:
    """Add new skills to the Technical Skills section.

    Accepts either a flat list of strings or a list of category groups
    ({"category": str, "skills": [str]}). Skips skills already present anywhere
    in the section (dedup) and places categorized skills onto the matching
    category line instead of dumping everything on the last one.
    """
    if not skills:
        return False

    heading_idx = _find_section(paragraphs, "TECHNICAL SKILLS", "SKILLS")
    if heading_idx is None:
        return False

    start, end = _section_content_range(paragraphs, heading_idx)
    skill_paras = [paragraphs[i] for i in range(start, end) if _para_text(paragraphs[i])]
    if not skill_paras:
        return False

    last_para = skill_paras[-1]
    existing_lower = " ".join(_para_text(p) for p in skill_paras).lower()
    seen = set()  # avoid re-adding within this call

    def _filter(items: list) -> list:
        out = []
        for s in items:
            s = (s or "").strip()
            key = s.lower()
            if not s or key in seen or key in existing_lower:
                continue
            seen.add(key)
            out.append(s)
        return out

    applied = False

    # Categorized form: [{"category": ..., "skills": [...]}, ...]
    if skills and isinstance(skills[0], dict):
        for group in skills:
            fresh = _filter(group.get("skills", []))
            if not fresh:
                continue
            target = _best_category_para(skill_paras, group.get("category", "")) or last_para
            _append_to_para(target, fresh)
            applied = True
        return applied

    # Flat list of strings: dedup, then append remaining to the last line.
    fresh = _filter(skills)
    if not fresh:
        return False
    _append_to_para(last_para, fresh)
    return True


_COMPANY_STOPWORDS = {
    "pvt", "private", "ltd", "limited", "inc", "llc", "llp", "co", "company",
    "corp", "corporation", "solutions", "technologies", "technology", "tech",
    "labs", "systems", "services", "group", "the", "and",
}


def _company_core_tokens(name: str) -> set:
    return {w for w in re.findall(r"[a-z0-9]+", name.lower())
            if len(w) > 1 and w not in _COMPANY_STOPWORDS}


def _company_matches(para_text: str, company: str) -> bool:
    """Match a job line to a company name, tolerant of suffix differences
    (e.g. 'Spirit AI Solutions' vs 'Spirit AI Solutions Pvt Ltd')."""
    pt = para_text.lower()
    if company.lower() in pt:
        return True
    core = _company_core_tokens(company)
    if not core:
        return False
    present = {t for t in core if re.search(rf"\b{re.escape(t)}\b", pt)}
    # Match when the meaningful name tokens are present (allow one to differ).
    return len(present) >= max(1, len(core) - 1)


def _add_experience_bullets(paragraphs, enhancements: list) -> int:
    """Add new bullet points to specified jobs in experience section. Returns count of jobs enhanced."""
    if not enhancements:
        return 0

    exp_heading_idx = _find_section(paragraphs, "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EXPERIENCE")
    if exp_heading_idx is None:
        return 0

    exp_start, exp_end = _section_content_range(paragraphs, exp_heading_idx)
    exp_paras = paragraphs[exp_start:exp_end]

    applied = 0
    for enh in enhancements:
        company = enh.get("company", "").strip()
        new_bullets = enh.get("new_bullets", [])
        if not company or not new_bullets:
            continue

        # Find the job-title paragraph for this company (suffix-tolerant).
        job_para_idx = None
        for i, p in enumerate(exp_paras):
            if _company_matches(_para_text(p), company):
                job_para_idx = i
                break

        if job_para_idx is None:
            continue

        # Find the last real bullet paragraph for this job, stopping at the next
        # job heading ("Title | Company   Date").
        last_bullet_para = None
        for i in range(job_para_idx + 1, len(exp_paras)):
            p = exp_paras[i]
            text = _para_text(p)
            if not text:
                continue
            if "|" in text and not _is_bullet_para(p):
                break  # next job's title line
            if _is_bullet_para(p):
                last_bullet_para = p

        # No real bullet found — skip rather than dumping bold lines at the title.
        # (Inserting at the title is what produced the bold, mis-positioned bullets.)
        if last_bullet_para is None:
            continue

        # Match existing bullet style: if the list auto-renders a glyph (numbering
        # or a List style), don't add a literal one; otherwise mirror the glyph used.
        if _auto_bulleted(last_bullet_para):
            prefix = ""
        else:
            glyph = _leading_glyph(_para_text(last_bullet_para)) or "●"
            prefix = f"{glyph} "

        # Insert new bullets after the last existing bullet (reversed keeps order).
        for bullet_text in reversed(new_bullets):
            _insert_para_after(last_bullet_para, f"{prefix}{bullet_text}")
        applied += 1

    return applied


def apply_edits_to_docx(original_bytes: bytes, edits: dict) -> tuple[bytes, dict]:
    """Apply AI-suggested edits to original DOCX preserving full structure.

    Returns (docx_bytes, report). report has booleans/counts so callers can
    detect when section heuristics failed to match anything (silent no-op).
    """
    doc = Document(io.BytesIO(original_bytes))
    paragraphs = _get_all_paragraphs(doc)

    report = {"summary_updated": False, "skills_added": False, "jobs_enhanced": 0}

    if edits.get("updated_summary"):
        report["summary_updated"] = _update_summary(paragraphs, edits["updated_summary"])

    if edits.get("skills_to_add"):
        report["skills_added"] = _add_skills(paragraphs, edits["skills_to_add"])

    if edits.get("experience_enhancements"):
        report["jobs_enhanced"] = _add_experience_bullets(paragraphs, edits["experience_enhancements"])

    report["any_applied"] = (
        report["summary_updated"] or report["skills_added"] or report["jobs_enhanced"] > 0
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), report
